import os
import json
import shutil
from typing import List, Dict, Any, Optional
from sqlalchemy.orm import Session
from docxtpl import DocxTemplate
from app.models.template import Template
from app.models.folder import Folder
from app.core.config import settings

class TemplateService:
    def __init__(self, db: Session):
        self.db = db
        self.templates_dir = settings.TEMPLATES_DIR

    def _get_template_file_path(self, template: Template) -> str:
        """Получает путь к файлу шаблона"""
        return os.path.join(self.templates_dir, template.filename)

    def upload_template(self, file, folder_id: int, user_id: int) -> Template:
        """Загружает шаблон в указанную папку"""
        # Создаем папку для шаблонов, если её нет
        os.makedirs(self.templates_dir, exist_ok=True)
        
        # Сохраняем файл
        filename = file.filename
        file_path = os.path.join(self.templates_dir, filename)
        
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Создаем запись в БД
        template = Template(
            filename=filename,
            folder_id=folder_id,
            uploaded_by=user_id
        )
        
        self.db.add(template)
        self.db.commit()
        self.db.refresh(template)
        
        return template

    def get_templates_by_folder(self, folder_id: int) -> List[Template]:
        """Получает все шаблоны в папке"""
        return self.db.query(Template).filter(Template.folder_id == folder_id).all()

    def get_template_by_id(self, template_id: int) -> Optional[Template]:
        """Получает шаблон по ID"""
        return self.db.query(Template).filter(Template.id == template_id).first()

    def delete_template(self, template_id: int) -> bool:
        """Удаляет шаблон"""
        template = self.get_template_by_id(template_id)
        if not template:
            return False
        
        # Удаляем файл
        file_path = self._get_template_file_path(template)
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Удаляем запись из БД
        self.db.delete(template)
        self.db.commit()
        
        return True

    def extract_placeholders(self, template_id: int) -> List[str]:
        """Извлекает плейсхолдеры из шаблона используя docxtpl"""
        template = self.get_template_by_id(template_id)
        if not template:
            print(f"Шаблон с ID {template_id} не найден")
            return []
        
        print(f"Извлекаем плейсхолдеры из шаблона: {template.filename}")
        file_path = self._get_template_file_path(template)
        print(f"Путь к файлу: {file_path}")
        
        # Проверяем существование файла
        if not os.path.exists(file_path):
            print(f"Файл не существует: {file_path}")
            return []
        
        try:
            # Используем python-docx напрямую для извлечения плейсхолдеров
            from docx import Document
            doc = Document(file_path)
            
            # Получаем все плейсхолдеры из шаблона
            placeholders = []
            import re
            
            # Извлекаем плейсхолдеры из параграфов
            for paragraph in doc.paragraphs:
                text = paragraph.text
                # Ищем плейсхолдеры в формате {{ любой текст }}
                matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
                placeholders.extend(matches)
            
            # Извлекаем плейсхолдеры из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
                            placeholders.extend(matches)
            
            # Также извлекаем плейсхолдеры из заголовков и футеров
            for section in doc.sections:
                if section.header:
                    for header in section.header.paragraphs:
                        text = header.text
                        matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
                        placeholders.extend(matches)
                
                if section.footer:
                    for footer in section.footer.paragraphs:
                        text = footer.text
                        matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
                        placeholders.extend(matches)
            
            # Очищаем плейсхолдеры от лишних пробелов и сохраняем порядок появления
            cleaned_placeholders = []
            for placeholder in placeholders:
                cleaned = placeholder.strip()
                if cleaned and cleaned not in cleaned_placeholders:
                    cleaned_placeholders.append(cleaned)
            
            # Возвращаем в порядке появления в документе (без сортировки)
            print(f"Найдены плейсхолдеры: {cleaned_placeholders}")
            return cleaned_placeholders
            
        except Exception as e:
            print(f"Ошибка извлечения плейсхолдеров: {e}")
            return []

    def generate_document(self, template_id: int, values: Dict[str, Any], output_format: str = 'docx') -> str:
        """Генерирует документ с заменой плейсхолдеров используя docxtpl"""
        template = self.get_template_by_id(template_id)
        if not template:
            raise ValueError("Шаблон не найден")
        
        try:
            from docxtpl import DocxTemplate
            from jinja2 import Environment, BaseLoader
            import tempfile
            import re
            
            # Сначала проверяем и исправляем незакрытые плейсхолдеры в шаблоне
            template_path = self._get_template_file_path(template)
            try:
                from docx import Document
                template_doc = Document(template_path)
                fixed = False
                
                # Проверяем и исправляем каждый параграф
                for paragraph in template_doc.paragraphs:
                    text = paragraph.text
                    original_text = text
                    # Ищем незакрытые плейсхолдеры вида {{переменная} (без второй закрывающей скобки)
                    # Паттерн: {{ что-то } но не }}
                    matches = list(re.finditer(r'\{\{([^}]+)\}(?!\})', text))
                    if matches:
                        # Исправляем с конца, чтобы индексы не сбивались
                        for match in reversed(matches):
                            placeholder_name = match.group(1).strip()
                            old_text = match.group(0)
                            new_text = f"{{{{{placeholder_name}}}}}"
                            text = text[:match.start()] + new_text + text[match.end():]
                            fixed = True
                            print(f"Исправлен незакрытый плейсхолдер: {old_text} -> {new_text}")
                    
                    if fixed and text != original_text:
                        paragraph.text = text
                
                # Также проверяем таблицы
                for table in template_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                text = paragraph.text
                                original_text = text
                                matches = list(re.finditer(r'\{\{([^}]+)\}(?!\})', text))
                                if matches:
                                    for match in reversed(matches):
                                        placeholder_name = match.group(1).strip()
                                        old_text = match.group(0)
                                        new_text = f"{{{{{placeholder_name}}}}}"
                                        text = text[:match.start()] + new_text + text[match.end():]
                                        fixed = True
                                        print(f"Исправлен незакрытый плейсхолдер в таблице: {old_text} -> {new_text}")
                                    if text != original_text:
                                        paragraph.text = text
                
                # Если были исправления, сохраняем временную копию
                if fixed:
                    temp_file = tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False)
                    temp_path = temp_file.name
                    temp_file.close()
                    template_doc.save(temp_path)
                    print(f"Шаблон исправлен и сохранен во временный файл: {temp_path}")
                    template_path = temp_path
                    # Удалим временный файл после использования
                    import atexit
                    atexit.register(lambda: os.remove(temp_path) if os.path.exists(temp_path) else None)
            except Exception as e:
                print(f"Не удалось проверить/исправить шаблон: {e}, используем оригинальный файл")
            
            # Создаем кастомный Jinja2 environment с более гибкими настройками
            jinja_env = Environment(
                loader=BaseLoader(),
                autoescape=False,  # Отключаем автоэкранирование для docx
                trim_blocks=True,
                lstrip_blocks=True,
                keep_trailing_newline=True
            )
            
            # Загружаем шаблон с помощью docxtpl и передаем кастомный environment
            doc = DocxTemplate(template_path, jinja_env=jinja_env)
            
            print(f"Генерируем документ с контекстом: {values}")
            
            # Рендерим шаблон с обработкой ошибок Jinja2
            try:
                doc.render(values)
            except Exception as render_error:
                error_msg = str(render_error)
                print(f"Ошибка рендеринга шаблона: {error_msg}")
                import traceback
                print(f"Traceback: {traceback.format_exc()}")
                
                # Проверяем шаблон на наличие проблемных мест
                try:
                    from docx import Document
                    template_doc = Document(self._get_template_file_path(template))
                    problematic_lines = []
                    for i, paragraph in enumerate(template_doc.paragraphs):
                        open_braces = paragraph.text.count('{{')
                        close_braces = paragraph.text.count('}}')
                        if open_braces != close_braces:
                            problematic_lines.append(f"Строка {i}: {paragraph.text[:100]}... ({{{{: {open_braces}, }}: {close_braces})")
                    if problematic_lines:
                        print("Проблемные места в шаблоне:")
                        for line in problematic_lines[:5]:
                            print(f"  {line}")
                except Exception as e:
                    print(f"Не удалось проверить шаблон: {e}")
                
                # Проверяем, есть ли проблемные значения
                for key, value in values.items():
                    if isinstance(value, str) and ('{' in value or '}' in value):
                        print(f"Предупреждение: значение '{key}' содержит фигурные скобки: {value}")
                
                # Если ошибка связана с синтаксисом шаблона, даем более понятное сообщение
                if "unexpected" in error_msg.lower() or "syntax" in error_msg.lower():
                    raise ValueError(f"Ошибка синтаксиса в шаблоне '{template.filename}': {error_msg}. Проверьте шаблон на наличие некорректного синтаксиса Jinja2. Убедитесь, что все плейсхолдеры имеют формат {{переменная}} и правильно закрыты.")
                else:
                    raise ValueError(f"Ошибка рендеринга шаблона: {error_msg}")
            
            # Создаем папку для сгенерированных файлов
            output_dir = os.path.join(self.templates_dir, 'generated')
            os.makedirs(output_dir, exist_ok=True)
            
            # Генерируем имя файла
            base_name = os.path.splitext(template.filename)[0]
            output_filename = f"generated_{base_name}.docx"
            output_path = os.path.join(output_dir, output_filename)
            
            # Сохраняем документ
            doc.save(output_path)
            
            print(f"Документ успешно сгенерирован: {output_path}")
            
            # Удаляем временный файл, если он был создан
            if template_path != self._get_template_file_path(template) and os.path.exists(template_path):
                try:
                    os.remove(template_path)
                    print(f"Временный файл удален: {template_path}")
                except Exception as e:
                    print(f"Не удалось удалить временный файл: {e}")
            
            # Конвертируем в PDF, если нужно
            if output_format == 'pdf':
                pdf_path = self._convert_to_pdf(output_path)
                return pdf_path
            
            return output_path
            
        except ValueError as ve:
            # Пробрасываем ValueError как есть
            raise ve
        except Exception as e:
            print(f"Ошибка генерации документа: {e}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")
            raise ValueError(f"Ошибка генерации документа: {str(e)}")

    def _replace_placeholders_in_paragraph(self, paragraph, values):
        """Заменяет плейсхолдеры в параграфе с сохранением форматирования"""
        try:
            # Проверяем, есть ли плейсхолдеры в параграфе
            paragraph_text = paragraph.text
            has_placeholder = False
            for key in values.keys():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph_text:
                    has_placeholder = True
                    break
            
            if not has_placeholder:
                return
            
            # Более точный подход: заменяем плейсхолдеры в каждом run отдельно
            for run in paragraph.runs:
                run_text = run.text
                for key, value in values.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in run_text:
                        run_text = run_text.replace(placeholder, str(value))
                
                # Обновляем текст run только если он изменился
                if run_text != run.text:
                    run.text = run_text
                
        except Exception as e:
            print(f"Ошибка замены плейсхолдеров в параграфе: {e}")
            # Fallback: простая замена
            for key, value in values.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

    def _convert_to_pdf(self, docx_path: str) -> str:
        """Конвертирует DOCX в PDF используя LibreOffice"""
        try:
            import subprocess
            import tempfile
            
            # Создаем временную папку для конвертации
            temp_dir = tempfile.mkdtemp()
            
            # Путь к LibreOffice
            libreoffice_path = "/usr/bin/libreoffice"
            
            # Команда для конвертации
            cmd = [
                libreoffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", temp_dir,
                docx_path
            ]
            
            # Выполняем конвертацию
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode != 0:
                print(f"Ошибка конвертации в PDF: {result.stderr}")
                raise ValueError("Ошибка конвертации в PDF")
            
            # Находим сгенерированный PDF файл
            pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
            pdf_path = os.path.join(temp_dir, pdf_filename)
            
            if not os.path.exists(pdf_path):
                raise ValueError("PDF файл не был создан")
            
            # Копируем PDF в папку с шаблонами
            final_pdf_path = os.path.join(self.templates_dir, 'generated', pdf_filename)
            shutil.copy2(pdf_path, final_pdf_path)
            
            # Очищаем временную папку
            shutil.rmtree(temp_dir)
            
            return final_pdf_path
            
        except Exception as e:
            print(f"Ошибка конвертации в PDF: {e}")
            raise ValueError(f"Ошибка конвертации в PDF: {str(e)}")

    def get_all_templates(self) -> List[Template]:
        """Получает все шаблоны"""
        return self.db.query(Template).all()

    def get_templates_by_user(self, user_id: int) -> List[Template]:
        """Получает шаблоны пользователя"""
        return self.db.query(Template).filter(Template.uploaded_by == user_id).all()
