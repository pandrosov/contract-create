from fastapi import FastAPI, File, UploadFile, Form, Depends
from fastapi.responses import JSONResponse
from typing import List
from docx import Document
import re
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import io
import os
import shutil
from fastapi import HTTPException
from urllib.parse import quote
from fastapi.security import OAuth2PasswordRequestForm
from sqlalchemy.orm import Session
from db import SessionLocal, User
from auth import (
    get_password_hash, verify_password, create_access_token, get_current_user, require_admin, get_user_by_username
)

app = FastAPI()

# Функция для поиска всех {{ПОЛЕ}} в docx
FIELD_PATTERN = re.compile(r"\{\{(.*?)\}\}")

def extract_fields_from_docx(file_bytes: bytes) -> List[str]:
    doc = Document(io.BytesIO(file_bytes))
    fields = set()
    for para in doc.paragraphs:
        fields.update(FIELD_PATTERN.findall(para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fields.update(FIELD_PATTERN.findall(cell.text))
    return list(fields)

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), 'templates')
os.makedirs(TEMPLATES_DIR, exist_ok=True)

@app.post("/upload-template")
async def upload_template(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        return JSONResponse(status_code=400, content={"error": "Только .docx файлы поддерживаются"})
    save_path = os.path.join(TEMPLATES_DIR, file.filename)
    if os.path.exists(save_path):
        return JSONResponse(status_code=409, content={"error": "Шаблон с таким именем уже существует"})
    with open(save_path, "wb") as f:
        f.write(await file.read())
    return {"status": "ok", "filename": file.filename}

@app.get("/templates")
def list_templates():
    files = [f for f in os.listdir(TEMPLATES_DIR) if f.endswith('.docx')]
    return {"templates": files}

@app.get("/template-fields/{filename}")
def get_template_fields(filename: str):
    file_path = os.path.join(TEMPLATES_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Шаблон не найден")
    with open(file_path, "rb") as f:
        file_bytes = f.read()
    fields = extract_fields_from_docx(file_bytes)
    return {"fields": fields}

@app.delete("/template/{filename}")
def delete_template(filename: str):
    file_path = os.path.join(TEMPLATES_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Шаблон не найден")
    os.remove(file_path)
    return {"status": "deleted", "filename": filename}

class GenerateRequest(BaseModel):
    values: dict

@app.post("/generate-document")
async def generate_document(filename: str = Form(...), values: str = Form(...)):
    """
    filename: имя шаблона docx
    values: JSON-строка с подстановками
    """
    import json
    file_path = os.path.join(TEMPLATES_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Шаблон не найден")
    with open(file_path, "rb") as f:
        file_bytes = f.read()
    values_dict = json.loads(values)
    doc = Document(io.BytesIO(file_bytes))
    # Заменяем поля в параграфах
    for para in doc.paragraphs:
        for key, val in values_dict.items():
            para.text = para.text.replace(f"{{{{{key}}}}}", str(val))
    # Заменяем поля в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in values_dict.items():
                    cell.text = cell.text.replace(f"{{{{{key}}}}}", str(val))
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    # Корректное имя файла для Content-Disposition
    safe_ascii = f"generated_{filename}".encode('ascii', 'ignore').decode('ascii') or 'contract.docx'
    safe_utf8 = quote(f"generated_{filename}")
    headers = {
        "Content-Disposition": f"attachment; filename={safe_ascii}; filename*=UTF-8''{safe_utf8}"
    }
    return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)

# Зависимость для получения сессии БД
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

class UserCreate(BaseModel):
    username: str
    email: str
    password: str

class UserOut(BaseModel):
    id: int
    username: str
    email: str
    is_active: bool
    is_admin: bool
    class Config:
        from_attributes = True

# Регистрация пользователя
@app.post("/register", response_model=UserOut)
def register(user: UserCreate, db: Session = Depends(get_db)):
    if db.query(User).filter((User.username == user.username) | (User.email == user.email)).first():
        raise HTTPException(status_code=400, detail="Пользователь с таким именем или email уже существует")
    db_user = User(
        username=user.username,
        email=user.email,
        password_hash=get_password_hash(user.password),
        is_active=False,
        is_admin=False
    )
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    return db_user

# Вход (получение токена)
@app.post("/token")
def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = get_user_by_username(db, form_data.username)
    if not user or not verify_password(form_data.password, user.password_hash):
        raise HTTPException(status_code=400, detail="Неверное имя пользователя или пароль")
    if not user.is_active:
        raise HTTPException(status_code=403, detail="Пользователь не активирован администратором")
    access_token = create_access_token(data={"sub": user.username})
    return {"access_token": access_token, "token_type": "bearer"}

# Получение информации о себе
@app.get("/me", response_model=UserOut)
def get_me(current_user: User = Depends(get_current_user)):
    return current_user

# Активация пользователя (только для админа)
class ActivateUserRequest(BaseModel):
    user_id: int
    is_active: bool

@app.post("/activate-user")
def activate_user(req: ActivateUserRequest, db: Session = Depends(get_db), admin: User = Depends(require_admin)):
    user = db.query(User).filter(User.id == req.user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    user.is_active = req.is_active
    db.commit()
    return {"status": "ok", "user_id": user.id, "is_active": user.is_active} 